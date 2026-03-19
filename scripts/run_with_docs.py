#!/usr/bin/env python3
"""
Run ReChunk on a set of documents: build index from cache, then query.

Chunking is done by the Temporal worker only. Run the worker and scripts/start_strategy_chunking.py
to backfill chunks for each strategy; this script only reads from storage/strategies/ and never runs
chunking (LLM or built-in) itself.

Usage:
  python scripts/run_with_docs.py <path>                    # index from cache, then remind to use --query
  python scripts/run_with_docs.py <path> --query "?"         # one query with retrieval + LLM feedback
  python scripts/run_with_docs.py <path> --interactive        # index from cache, then prompt for questions

Requires OPENAI_API_KEY for LLM synthesis at query time.
"""

import argparse
import asyncio
import hashlib
import json
import sys
import time
from dataclasses import dataclass
from pathlib import Path

# Project root and src on path (rechunk + temporal_workflows)
_script_dir = Path(__file__).resolve().parent
_project_root = _script_dir.parent
sys.path.insert(0, str(_project_root / "src"))
sys.path.insert(0, str(_project_root))

from llama_index.core import Document, VectorStoreIndex, Settings
from llama_index.core.schema import MetadataMode, QueryBundle, TextNode
from llama_index.llms.openai import OpenAI
from llama_index.embeddings.openai import OpenAIEmbedding

from rechunk.cache import cache_updated_since, get_strategy_cache_mtimes

# Max characters of chunk text to show in retrieval feedback
CHUNK_PREVIEW_LEN = 280


@dataclass
class Strategy:
    """In-memory representation of a chunking strategy."""

    id: str
    kind: str  # "builtin_splitter" (no LLM) or "llm"
    instruction: str
    # For kind="builtin_splitter": which LlamaIndex parser. "sentence" | "token"
    splitter: str = "sentence"
    # For kind="llm": OpenAI model name (default gpt-4o-mini when None)
    model: str | None = None


# Strategies metadata file: project root (next to pyproject.toml)
STRATEGIES_FILE = Path(__file__).resolve().parent.parent / "rechunk_strategies.json"
# Chunk cache directory for per-strategy, per-doc LLM outputs. This is a
# pragmatic first version; we may later move to a more structured storage
# layout or a vector-store-native persistence strategy.
STRATEGY_CACHE_DIR = Path(__file__).resolve().parent.parent / "storage" / "strategies"

# Default when no strategy file exists: baseline (built-in) only, never LLM.
DEFAULT_BASELINE_STRATEGY = Strategy(
    id="s_default",
    kind="builtin_splitter",
    instruction="Sentence-based splitting (LlamaIndex default, chunk_size=1024)",
    splitter="sentence",
)


def _strategy_to_dict(s: Strategy) -> dict:
    return {
        "id": s.id,
        "kind": s.kind,
        "instruction": s.instruction,
        "splitter": getattr(s, "splitter", "sentence"),
        "model": getattr(s, "model", None),
    }


def _dict_to_strategy(d: dict) -> Strategy:
    return Strategy(
        id=d["id"],
        kind=d["kind"],
        instruction=d["instruction"],
        splitter=d.get("splitter", "sentence"),
        model=d.get("model"),
    )


def load_strategies(path: Path) -> list[Strategy] | None:
    """Load strategy set from JSON file. Returns None if file missing or invalid."""
    if not path.exists():
        return None
    try:
        raw = path.read_text(encoding="utf-8")
        data = json.loads(raw)
        if not isinstance(data, list) or len(data) == 0:
            return None
        return [_dict_to_strategy(item) for item in data]
    except (json.JSONDecodeError, KeyError, TypeError):
        return None


def save_strategies(path: Path, strategies: list[Strategy]) -> None:
    """Serialize strategy set to JSON."""
    data = [_strategy_to_dict(s) for s in strategies]
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def _strategy_cache_path(strategy_id: str) -> Path:
    return STRATEGY_CACHE_DIR / f"{strategy_id}_chunks.jsonl"


def _node_to_dict(node: TextNode) -> dict:
    if hasattr(node, "get_content"):
        text = node.get_content(metadata_mode=MetadataMode.NONE)
    else:
        text = getattr(node, "text", "")
    return {
        "id": getattr(node, "id_", None) or "",
        "text": text,
        "metadata": getattr(node, "metadata", None) or {},
        "ref_doc_id": getattr(node, "ref_doc_id", None),
        "start_char_idx": getattr(node, "start_char_idx", None),
        "end_char_idx": getattr(node, "end_char_idx", None),
    }


def _dict_to_node(d: dict) -> TextNode:
    node = TextNode(
        id_=d.get("id", ""),
        text=d.get("text", ""),
        metadata=d.get("metadata") or {},
        ref_doc_id=d.get("ref_doc_id"),
    )
    start = d.get("start_char_idx")
    end = d.get("end_char_idx")
    if start is not None and end is not None:
        node.start_char_idx = start
        node.end_char_idx = end
    return node


def _load_strategy_chunk_cache(strategy_id: str) -> dict[str, list[TextNode]]:
    """
    Load cached chunks for a strategy, keyed by document content_hash.

    Cache format is deliberately simple JSONL for now; we may later migrate to a
    more structured store or a proper vector-store-backed cache.
    """
    path = _strategy_cache_path(strategy_id)
    if not path.exists():
        return {}
    cache: dict[str, list[TextNode]] = {}
    try:
        with path.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                    h = rec.get("content_hash")
                    nodes_data = rec.get("nodes") or []
                    if not h or not isinstance(nodes_data, list):
                        continue
                    cache[h] = [_dict_to_node(n) for n in nodes_data]
                except json.JSONDecodeError:
                    continue
    except OSError:
        return {}
    return cache


def _append_strategy_chunk_cache(strategy_id: str, content_hash: str, nodes: list[TextNode]) -> None:
    """
    Append per-doc chunk results to the per-strategy JSONL cache.

    This is intentionally append-only and JSONL-based for simplicity; we may
    later revise it to support compaction or a different storage backend.
    """
    STRATEGY_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    path = _strategy_cache_path(strategy_id)
    rec = {
        "content_hash": content_hash,
        "nodes": [_node_to_dict(n) for n in nodes],
    }
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(rec) + "\n")


def _extract_file_content(path: Path) -> tuple[str | None, str]:
    """Extract text content from supported file types.

    Returns (content, description). content is None if the file type isn't supported
    or required libraries are missing.
    """
    import os

    ext = path.suffix.lower()
    file_size = path.stat().st_size

    try:
        if ext == ".pdf":
            try:
                import PyPDF2  # type: ignore[import]

                with path.open("rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text_parts: list[str] = []
                    for page in reader.pages:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
                content = "\n".join(text_parts)
                return content, f"PDF document ({file_size:,} bytes)"
            except ImportError:
                print(f"[WARN] PyPDF2 not installed; skipping PDF {path}", file=sys.stderr)
                return None, "PDF document (PyPDF2 not installed)"

        if ext == ".docx":
            try:
                import docx  # type: ignore[import]

                doc = docx.Document(str(path))
                paragraphs = [p.text for p in doc.paragraphs]
                content = "\n".join(paragraphs)
                return content, f"Word document ({file_size:,} bytes)"
            except ImportError:
                print(f"[WARN] python-docx not installed; skipping DOCX {path}", file=sys.stderr)
                return None, "Word document (python-docx not installed)"

        # Plain text-like files
        if ext in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8", errors="replace")
            return text, f"Text file ({ext})"

        # Fallback: try as text
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            if text.strip():
                return text, f"Unknown text file ({ext})"
        except Exception:
            pass

        return None, f"Unsupported or binary file ({ext})"
    except Exception as e:
        return None, f"Error reading {path}: {e}"


def _content_hash(content: str) -> str:
    """SHA-256 hash of document content for deduplication."""
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def _split_long_nodes(
    nodes: list,
    *,
    # OpenAI embedding endpoint hard-limit is 8192 tokens for some models; use a safety margin.
    max_embed_tokens: int = 7500,
    overlap_tokens: int = 200,
    max_chars_fallback: int = 12000,
) -> list:
    """
    Ensure no node content exceeds embedding limits by splitting large nodes.

    This is a pragmatic guardrail for embedding models with hard context limits
    (e.g., 8k tokens). We may later replace this with a more principled,
    token-aware splitter or embedder-specific policy.
    """
    # Token-aware splitting using tiktoken (matches OpenAI tokenization more closely than char counts).
    enc = None
    try:
        import tiktoken  # type: ignore

        try:
            enc = tiktoken.encoding_for_model("text-embedding-3-small")
        except Exception:
            enc = tiktoken.get_encoding("cl100k_base")
    except Exception:
        enc = None

    result: list[TextNode] = []
    for node in nodes:
        # Prefer get_content when available; otherwise, fall back to .text
        if hasattr(node, "get_content"):
            text = node.get_content(metadata_mode=MetadataMode.NONE)
        else:
            text = getattr(node, "text", "")
        if not isinstance(text, str):
            text = str(text)

        # Fast path: token count under limit (or char fallback under limit).
        token_count = None
        if enc is not None:
            try:
                token_count = len(enc.encode(text))
            except Exception:
                token_count = None
        if (token_count is not None and token_count <= max_embed_tokens) or (
            token_count is None and len(text) <= max_chars_fallback
        ):
            result.append(node)
            continue

        # Split oversized node into smaller windows.
        meta = getattr(node, "metadata", None) or {}
        ref_doc_id = getattr(node, "ref_doc_id", None) or meta.get("source_doc", "")
        base_id = getattr(node, "id_", None) or ref_doc_id or "node"
        part_idx = 1
        if enc is not None:
            try:
                toks = enc.encode(text)
                step = max(1, max_embed_tokens - max(0, overlap_tokens))
                start = 0
                while start < len(toks):
                    window = toks[start : start + max_embed_tokens]
                    chunk_text = enc.decode(window)
                    part_id = f"{base_id}_part{part_idx}"
                    result.append(TextNode(id_=part_id, text=chunk_text, metadata=dict(meta), ref_doc_id=ref_doc_id))
                    part_idx += 1
                    start += step
                continue
            except Exception:
                # Fall back to char splitting below.
                pass

        # Char-based fallback split (best-effort).
        start = 0
        while start < len(text):
            chunk_text = text[start : start + max_chars_fallback]
            part_id = f"{base_id}_part{part_idx}"
            result.append(TextNode(id_=part_id, text=chunk_text, metadata=dict(meta), ref_doc_id=ref_doc_id))
            start += max_chars_fallback
            part_idx += 1
    return result


def load_documents(path: Path) -> list[Document]:
    """Load supported files from a file or directory into LlamaIndex Documents.

    For a directory, scans the entire directory tree recursively (all subdirectories).
    Supports .txt, .md, .pdf, .docx. Other text-like files are best-effort.
    Documents with identical content (same hash) are loaded only once.
    """
    path = path.resolve()
    if not path.exists():
        raise FileNotFoundError(f"Path not found: {path}")

    docs: list[Document] = []
    seen_hashes: set[str] = set()

    if path.is_file():
        content, desc = _extract_file_content(path)
        if content:
            h = _content_hash(content)
            docs.append(
                Document(
                    text=content,
                    id_=path.name,
                    metadata={"content_hash": h},
                )
            )
        else:
            raise FileNotFoundError(f"Unable to read file {path}: {desc}")
    else:
        # Recursively scan whole tree (all subdirectories)
        candidates = sorted(path.rglob("*"))
        for f in (p for p in candidates if p.is_file() and p.suffix.lower() in {".txt", ".md", ".pdf", ".docx"}):
            content, desc = _extract_file_content(f)
            if content:
                h = _content_hash(content)
                if h in seen_hashes:
                    print(f"[SKIP] Duplicate content (same hash): {f.relative_to(path)}", file=sys.stderr)
                    continue
                seen_hashes.add(h)
                doc_id = str(f.relative_to(path))
                docs.append(
                    Document(
                        text=content,
                        id_=doc_id,
                        metadata={"content_hash": h},
                    )
                )
            else:
                print(f"[WARN] Skipping {f}: {desc}", file=sys.stderr)
        if not docs:
            raise FileNotFoundError(f"No supported files under {path}")

    return docs


def build_index_for_strategies(
    strategies: list[Strategy],
    docs: list[Document],
    *,
    quiet: bool = False,
) -> tuple[VectorStoreIndex, list]:
    """Build index from cache only. All chunking is done by the Temporal worker; CLI never runs chunking."""
    from llama_index.core.schema import BaseNode

    all_nodes: list[BaseNode] = []
    n_docs = len(docs)
    for idx, s in enumerate(strategies, 1):
        if not quiet:
            print(f"\n  Strategy {idx}/{len(strategies)}: {s.id} ({s.kind}) — {n_docs} documents")
        cache = _load_strategy_chunk_cache(s.id)
        strat_nodes: list[BaseNode] = []
        for j, doc in enumerate(docs, 1):
            content_hash = (doc.metadata or {}).get("content_hash") if hasattr(doc, "metadata") else None
            if content_hash and content_hash in cache:
                strat_nodes.extend(cache[content_hash])
                if not quiet:
                    print(f"    [{j}/{n_docs}] cache hit for hash={content_hash[:12]}... ({doc.id_})")
            else:
                if not quiet:
                    print(
                        f"    [{j}/{n_docs}] no cached chunks for {getattr(doc, 'id_', '?')!r} (strategy {s.id}); "
                        "run worker + start_strategy_chunking to backfill."
                    )
        all_nodes.extend(strat_nodes)
        if not quiet:
            print(f"    → {len(strat_nodes)} chunks from cache")
    if not quiet:
        print(f"\n  Total before max-length enforcement: {len(all_nodes)} chunks from {len(strategies)} strategy(ies)")

    # Enforce a hard maximum text length per chunk before embedding. This is a
    # defensive measure against embedding API limits (e.g., 8k-token inputs).
    # We may later revise this to be token-aware or embedder-specific.
    all_nodes = _split_long_nodes(all_nodes)
    if not quiet:
        print(f"  Total after max-length enforcement: {len(all_nodes)} chunks")

    try:
        index = VectorStoreIndex(all_nodes)
    except Exception as e:
        # Embedding calls can still fail (rate limits, token limits, network). Provide a clearer message.
        raise RuntimeError(
            "Failed to build VectorStoreIndex (embedding step). "
            "This usually means one chunk exceeded the embed token limit or the embedding API failed. "
            "Try rerunning, or reduce chunk sizes."
        ) from e
    return index, all_nodes


def _trigger_strategy_chunking_sync(docs_root: Path, doc_ids: list[str], strategy: Strategy) -> None:
    """Start Temporal workflow for this strategy (so worker backfills chunks). Runs async client in sync context."""
    try:
        from temporalio.client import Client
        from temporal_workflows import StrategyChunkingInput, StrategyChunkingWorkflow
    except ImportError:
        print("  (Temporal not available; run scripts/start_strategy_chunking.py manually to backfill.)", file=sys.stderr)
        return
    TASK_QUEUE = "rechunk-strategy-chunking"

    async def _run() -> None:
        client = await Client.connect("localhost:7233")
        workflow_id = f"rechunk-{strategy.id}"
        await client.start_workflow(
            StrategyChunkingWorkflow,
            StrategyChunkingInput(
                strategy_id=strategy.id,
                kind=strategy.kind,
                docs_root=str(docs_root.resolve()),
                doc_ids=doc_ids,
                strategy_instruction=strategy.instruction if strategy.kind == "llm" else None,
                model=getattr(strategy, "model", None),
                splitter=getattr(strategy, "splitter", "sentence"),
            ),
            id=workflow_id,
            task_queue=TASK_QUEUE,
        )
        print(f"  Started workflow {workflow_id}. Worker will backfill chunks for strategy {strategy.id!r}.", file=sys.stderr)

    try:
        asyncio.run(_run())
    except Exception as e:
        print(f"  Could not start workflow: {e}. Run scripts/start_strategy_chunking.py manually if needed.", file=sys.stderr)


def manage_strategies_interactively(
    strategies: list[Strategy],
    strategies_path: Path,
    docs_root: Path | None = None,
    doc_ids: list[str] | None = None,
) -> None:
    """Simple CLI to add/remove strategies (built-in splitters or LLM). Saves after each add/delete.
    If docs_root and doc_ids are provided, starting a new strategy also triggers the Temporal workflow."""
    while True:
        print("\nCurrent strategies:")
        for i, s in enumerate(strategies, 1):
            extra = f"  splitter={s.splitter}" if s.kind == "builtin_splitter" else ""
            if s.kind == "llm":
                extra = f"  model={getattr(s, 'model', None) or 'gpt-4o-mini'}"
            print(f"  [{i}] {s.id}  (type={s.kind}{extra})  instruction={s.instruction!r}")
        print(
            "\nStrategy menu:\n"
            "  [a] Add new LLM strategy\n"
            "  [b] Add built-in splitter (Sentence or Token)\n"
            "  [d] Delete a strategy\n"
            "  [x] Back (no more changes)\n"
        )
        choice = input("Choice [a/b/d/x]: ").strip().lower()
        if choice in ("x", "", "q", "exit"):
            break
        if choice == "a":
            sid = input("New strategy id (e.g. s_procedures): ").strip()
            if not sid:
                print("No id entered; skipping.")
                continue
            instr = input("Strategy instruction (what semantic unit to chunk around): ").strip()
            if not instr:
                print("No instruction entered; skipping.")
                continue
            model_input = input("Model (default gpt-4o-mini, or e.g. gpt-4o): ").strip()
            model = model_input if model_input else None
            strategies.append(Strategy(id=sid, kind="llm", instruction=instr, model=model))
            save_strategies(strategies_path, strategies)
            print(f"Added LLM strategy {sid!r} (model={model or 'gpt-4o-mini'}). Saved to {strategies_path.name}")
            if docs_root is not None and doc_ids:
                _trigger_strategy_chunking_sync(docs_root, doc_ids, strategies[-1])
        elif choice == "b":
            which = input("Built-in splitter: [1] SentenceSplitter  [2] TokenTextSplitter  (1 or 2): ").strip()
            if which == "2":
                sid = "s_token"
                splitter = "token"
                instr = "Token-based splitting (LlamaIndex default chunk_size=1024, overlap=20)"
            else:
                sid = "s_sentence"
                splitter = "sentence"
                instr = "Sentence-based splitting (LlamaIndex default chunk_size=1024, overlap=20)"
            strategies.append(
                Strategy(id=sid, kind="builtin_splitter", instruction=instr, splitter=splitter),
            )
            save_strategies(strategies_path, strategies)
            print(f"Added built-in splitter {sid!r} ({splitter}). Saved to {strategies_path.name}")
            if docs_root is not None and doc_ids:
                _trigger_strategy_chunking_sync(docs_root, doc_ids, strategies[-1])
        elif choice == "d":
            idx_str = input("Enter number of strategy to delete (or blank to cancel): ").strip()
            if not idx_str:
                continue
            try:
                idx = int(idx_str)
            except ValueError:
                print("Invalid number.")
                continue
            if not (1 <= idx <= len(strategies)):
                print("Out of range.")
                continue
            removed = strategies.pop(idx - 1)
            save_strategies(strategies_path, strategies)
            print(f"Removed strategy {removed.id!r}. Saved to {strategies_path.name}")
        else:
            print("Unknown choice; please enter a/b/d/x.")


def run_query_with_feedback(
    index: VectorStoreIndex,
    query: str,
    top_k: int = 5,
    total_chunks: int | None = None,
    strategy_ids: list[str] | None = None,
) -> None:
    """
    Run retrieval (embedding comparison), show chunks + scores, then LLM synthesis.
    Prints timing and explains cosine-similarity retrieval vs LLM "dress-up".
    When total_chunks/strategy_ids are provided, shows that retrieval is over all strategies.
    """
    retriever = index.as_retriever(similarity_top_k=top_k)
    query_bundle = QueryBundle(query_str=query)

    # ---- Retrieval (embedding comparison) ----
    print("\n" + "=" * 60)
    print("RETRIEVAL (embedding comparison)")
    print("=" * 60)
    print(
        "Your question is embedded, then compared to each chunk's embedding via"
        " cosine similarity (higher score = more similar). No LLM call yet."
    )
    if total_chunks is not None and strategy_ids:
        print(f"  Pool: {total_chunks} chunks from all strategies combined: {', '.join(strategy_ids)}")
    t0 = time.perf_counter()
    nodes_with_scores = retriever.retrieve(query)
    t1 = time.perf_counter()
    print(f"  Time: {(t1 - t0) * 1000:.0f} ms")
    print(f"  Top {len(nodes_with_scores)} chunks:")
    for i, nws in enumerate(nodes_with_scores, 1):
        node = nws.node
        score = getattr(nws, "score", None)
        score_str = f"  score={score:.4f}" if score is not None else "  (no score)"
        if hasattr(node, "get_content"):
            text = node.get_content(metadata_mode=MetadataMode.NONE)
        else:
            text = getattr(node, "text", str(node))
        preview = (text[:CHUNK_PREVIEW_LEN] + "…") if len(text) > CHUNK_PREVIEW_LEN else text
        meta = getattr(node, "metadata", None) or {}
        source = getattr(node, "ref_doc_id", None) or meta.get("source_doc", "?")
        strategy = meta.get("strategy", "?")
        start = getattr(node, "start_char_idx", None)
        end = getattr(node, "end_char_idx", None)
        if start is not None and end is not None:
            loc = f"  source={source!r}  strategy={strategy}  chars {start}–{end}"
        else:
            loc = f"  source={source!r}  strategy={strategy}  (position in doc not stored)"
        print(f"    {i}.{score_str}{loc}")
        print(f"       {preview!r}")
    print()

    # ---- LLM synthesis ----
    print("=" * 60)
    print("LLM RESPONSE (synthesis from retrieved chunks)")
    print("=" * 60)
    print("The LLM sees your question + the chunks above and produces an answer.")
    engine = index.as_query_engine()
    t2 = time.perf_counter()
    response = engine.synthesize(query_bundle, nodes_with_scores)
    t3 = time.perf_counter()
    print(f"  Time: {(t3 - t2) * 1000:.0f} ms")
    answer = getattr(response, "response", str(response))
    print(f"  Answer: {answer}")
    print()


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Run ReChunk on documents: chunk with LLM strategy, build index, optional query."
    )
    parser.add_argument(
        "path",
        type=Path,
        help="Path to a .txt file or directory containing .txt files",
    )
    parser.add_argument(
        "--strategy-id",
        default="s_sections",
        help="Strategy id for optional LLM strategy (used only with --strategy)",
    )
    parser.add_argument(
        "--strategy",
        default=None,
        help="If set, add an LLM chunking strategy (instruction text). Default is no LLM; uses LlamaIndex SentenceSplitter only.",
    )
    parser.add_argument(
        "--query",
        default=None,
        help="If set, run this query against the index and print the response",
    )
    parser.add_argument(
        "--interactive",
        action="store_true",
        help="After chunking/indexing once, prompt for questions in a loop (fast retrieval feel)",
    )
    parser.add_argument(
        "--top-k",
        type=int,
        default=5,
        help="Number of chunks to retrieve per query (default: 5)",
    )
    parser.add_argument(
        "--model",
        default="gpt-4o-mini",
        help="OpenAI model for chunking (default: gpt-4o-mini)",
    )
    args = parser.parse_args()

    # Use OpenAI for LLM and embeddings (reads OPENAI_API_KEY from env)
    Settings.llm = OpenAI(model=args.model, temperature=0.1)
    Settings.embed_model = OpenAIEmbedding(model="text-embedding-3-small")

    docs = load_documents(args.path)
    print(f"Loaded {len(docs)} document(s) from {args.path}")

    # Load saved strategies if present; otherwise use default baseline and cue worker.
    strategies = load_strategies(STRATEGIES_FILE)
    if strategies is None:
        print()
        print("=" * 70)
        print("  WARNING: No strategy file found (or file empty/invalid).")
        print("  Using the default BASELINE strategy only (built-in splitter, NOT LLM).")
        print("  The worker will be cued to backfill chunks for this strategy.")
        print("  Ensure the Temporal worker is running (python temporal_worker.py).")
        print("=" * 70)
        print()
        strategies = [DEFAULT_BASELINE_STRATEGY]
        # Write a strategy file so the active strategy set is explicit/persistent.
        try:
            save_strategies(STRATEGIES_FILE, strategies)
            print(f"Wrote default strategies to {STRATEGIES_FILE.resolve()}")
            if not STRATEGIES_FILE.exists():
                print(
                    f"[WARN] Tried to write {STRATEGIES_FILE.resolve()} but it still does not exist.",
                    file=sys.stderr,
                )
        except Exception as e:
            print(f"[WARN] Could not write {STRATEGIES_FILE.resolve()}: {e}", file=sys.stderr)
        if args.strategy:
            strategies.append(
                Strategy(id=args.strategy_id, kind="llm", instruction=args.strategy),
            )
        # Cue the worker to backfill the default baseline so the user gets chunks.
        docs_root = Path(args.path).resolve().parent if Path(args.path).resolve().is_file() else Path(args.path).resolve()
        doc_ids = [getattr(d, "id_", "") for d in docs]
        _trigger_strategy_chunking_sync(docs_root, doc_ids, DEFAULT_BASELINE_STRATEGY)
    else:
        print(f"Loaded {len(strategies)} strategy(ies) from {STRATEGIES_FILE.name}")
        if args.strategy:
            strategies.append(
                Strategy(id=args.strategy_id, kind="llm", instruction=args.strategy),
            )

    index, nodes = build_index_for_strategies(strategies, docs)
    print(f"ReChunk produced {len(nodes)} nodes across {len(strategies)} strategy(ies)")
    print("Index built (embeddings computed).")

    if args.interactive:
        last_cache_mtimes = get_strategy_cache_mtimes([s.id for s in strategies])
        print(
            f"\n{len(nodes)} chunks formed from {len(strategies)} strategy(ies). "
            "Enter a question (or 'quit'/'q' to exit). Index is rebuilt automatically when the worker updates the cache."
        )
        while True:
            # Before each prompt: if worker wrote new chunks, rebuild index from cache (no user action needed).
            if cache_updated_since([s.id for s in strategies], last_cache_mtimes):
                try:
                    index, nodes = build_index_for_strategies(strategies, docs, quiet=True)
                    last_cache_mtimes = get_strategy_cache_mtimes([s.id for s in strategies])
                    print(f"\n  [Cache updated by worker; index rebuilt with {len(nodes)} chunks.]")
                except Exception as e:
                    # Don't crash the session if embeddings fail; keep the previous index.
                    print(f"\n  [WARN] Cache updated but index rebuild failed: {e}", file=sys.stderr)
            try:
                q = input("\nYour question: ").strip()
            except (EOFError, KeyboardInterrupt):
                print("\nBye.")
                break
            if not q:
                # Empty input should just reprompt (don't exit).
                continue
            if q.lower() in ("quit", "q", "exit"):
                print("Bye.")
                break
            if q.lower() in ("reload", "r"):
                index, nodes = build_index_for_strategies(strategies, docs, quiet=True)
                last_cache_mtimes = get_strategy_cache_mtimes([s.id for s in strategies])
                print(f"Index rebuilt: {len(nodes)} chunks.")
                continue
            run_query_with_feedback(
                index, q, top_k=args.top_k,
                total_chunks=len(nodes), strategy_ids=[s.id for s in strategies],
            )
            # Ask for feedback and optionally adjust strategies
            fb = input(
                "Was this answer correct? [y]es / [n]o / [i]ncomplete / [s]kip: "
            ).strip().lower()
            if fb in ("n", "i"):
                manage_strategies_interactively(
                    strategies, STRATEGIES_FILE,
                    docs_root=Path(args.path).resolve(),
                    doc_ids=[getattr(d, "id_", "") for d in docs],
                )
                # Rebuild index after any strategy changes
                index, nodes = build_index_for_strategies(strategies, docs)
                print(
                    f"\nRebuilt index: {len(nodes)} chunks from {len(strategies)} strategy(ies)."
                )
    elif args.query:
        run_query_with_feedback(
            index, args.query, top_k=args.top_k,
            total_chunks=len(nodes), strategy_ids=[s.id for s in strategies],
        )
    else:
        print("\nNo --query given. Use --query \"your question\" or --interactive to ask questions.")


if __name__ == "__main__":
    main()
