"""
Step 0.1 & 0.2: Verify activity logic, cache round-trip, and activity run inside Temporal.
Step 1.5: Verify activity retries (one doc fails once, retry succeeds; other docs complete).

0.1: Code path and cache work without Temporal. 0.2: Real activity runs in Temporal's
ActivityEnvironment and writes to cache. 1.5: Workflow + worker; one activity fails
on first attempt, Temporal retries it; both docs end up in cache.
"""

import asyncio
import os
from pathlib import Path
from unittest.mock import patch

import pytest
from llama_index.core import Document
from llama_index.core.schema import TextNode
from temporalio.testing import ActivityEnvironment

from rechunk.cache import (
    append_chunk_cache,
    compute_content_hash,
    load_chunk_cache,
)


def test_compute_content_hash():
    h = compute_content_hash("hello world")
    assert isinstance(h, str)
    assert len(h) == 64
    assert h == compute_content_hash("hello world")
    assert h != compute_content_hash("hello world.")


def test_cache_round_trip(tmp_path):
    """Append nodes to cache, load back, assert they match."""
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        strategy_id = "test_strategy_round_trip"
        content_hash = "abc123"
        nodes = [
            TextNode(
                id_="n1",
                text="chunk one",
                metadata={"strategy": strategy_id, "source_doc": "doc1"},
                ref_doc_id="doc1",
            ),
            TextNode(
                id_="n2",
                text="chunk two",
                metadata={"strategy": strategy_id, "source_doc": "doc1"},
                ref_doc_id="doc1",
            ),
        ]
        append_chunk_cache(strategy_id, content_hash, nodes)
        loaded = load_chunk_cache(strategy_id)
        assert content_hash in loaded
        got = loaded[content_hash]
        assert len(got) == 2
        assert got[0].text == "chunk one"
        assert got[0].metadata.get("strategy") == strategy_id
        assert got[0].metadata.get("source_doc") == "doc1"
        assert got[1].text == "chunk two"
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)


def test_activity_logic_in_isolation(tmp_path):
    """
    Run the exact pipeline the activity will use: read doc from path → parse (mocked) → append cache → load.
    Proves the code path works before it lives inside a Temporal activity.
    """
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        doc_content = "First paragraph.\n\nSecond paragraph."
        doc_file = tmp_path / "doc1.txt"
        doc_file.write_text(doc_content, encoding="utf-8")
        doc_id = "doc1.txt"
        strategy_id = "test_activity_logic"
        content_hash = compute_content_hash(doc_content)

        # What the activity will do: read text, build Document, run parser, append cache.
        text = doc_file.read_text(encoding="utf-8")
        doc = Document(text=text, id_=doc_id)

        # Mock LLMNodeParser.get_nodes_from_documents so we don't call the real LLM.
        mock_nodes = [
            TextNode(
                id_="doc1_chunk_1",
                text="First paragraph.",
                metadata={"strategy": strategy_id, "source_doc": doc_id},
                ref_doc_id=doc_id,
            ),
            TextNode(
                id_="doc1_chunk_2",
                text="Second paragraph.",
                metadata={"strategy": strategy_id, "source_doc": doc_id},
                ref_doc_id=doc_id,
            ),
        ]

        with patch("rechunk.node_parser.LLMNodeParser.get_nodes_from_documents", return_value=mock_nodes):
            from rechunk import LLMNodeParser

            parser = LLMNodeParser(
                strategy_id=strategy_id,
                strategy_instruction="Split by paragraph.",
            )
            nodes = parser.get_nodes_from_documents([doc])

        assert len(nodes) == 2
        append_chunk_cache(strategy_id, content_hash, nodes)
        loaded = load_chunk_cache(strategy_id)
        assert content_hash in loaded
        got = loaded[content_hash]
        assert len(got) == 2
        assert got[0].text == "First paragraph."
        assert got[1].text == "Second paragraph."
        assert got[0].metadata.get("source_doc") == doc_id
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)


@pytest.mark.asyncio
async def test_activity_runs_inside_temporal(tmp_path):
    """
    Step 0.2: Run the real chunk_doc_with_strategy activity in Temporal's activity context.
    Verifies serialization, async execution, and that the activity writes to the cache.
    """
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        doc_content = "Section A.\n\nSection B."
        doc_file = tmp_path / "doc.txt"
        doc_file.write_text(doc_content, encoding="utf-8")
        doc_id = "doc.txt"
        strategy_id = "test_temporal_activity"
        content_hash = compute_content_hash(doc_content)

        mock_nodes = [
            TextNode(
                id_="doc_chunk_1",
                text="Section A.",
                metadata={"strategy": strategy_id, "source_doc": doc_id},
                ref_doc_id=doc_id,
            ),
            TextNode(
                id_="doc_chunk_2",
                text="Section B.",
                metadata={"strategy": strategy_id, "source_doc": doc_id},
                ref_doc_id=doc_id,
            ),
        ]

        from temporal_activities import ChunkDocInput, chunk_doc_with_strategy

        activity_input = ChunkDocInput(
            strategy_id=strategy_id,
            strategy_instruction="Split by section.",
            model=None,
            docs_root=str(tmp_path),
            doc_id=doc_id,
            content_hash=content_hash,
        )

        with patch("rechunk.node_parser.LLMNodeParser.get_nodes_from_documents", return_value=mock_nodes):
            env = ActivityEnvironment()
            await env.run(chunk_doc_with_strategy, activity_input)

        loaded = load_chunk_cache(strategy_id)
        assert content_hash in loaded
        got = loaded[content_hash]
        assert len(got) == 2
        assert got[0].text == "Section A."
        assert got[1].text == "Section B."
        assert got[0].metadata.get("source_doc") == doc_id
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)


@pytest.mark.asyncio
async def test_builtin_chunking_runs_inside_temporal(tmp_path):
    """
    Run the real chunk_doc_with_builtin_splitter activity in Temporal's activity context.
    Uses real LlamaIndex SentenceSplitter (no mock). Verifies baseline chunking works in Temporal.
    """
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        # Content that SentenceSplitter will split into multiple chunks (multiple sentences).
        doc_content = "First sentence here. Second sentence there. Third one for good measure."
        doc_file = tmp_path / "builtin_doc.txt"
        doc_file.write_text(doc_content, encoding="utf-8")
        doc_id = "builtin_doc.txt"
        strategy_id = "test_builtin_temporal"
        content_hash = compute_content_hash(doc_content)

        from temporal_activities import BuiltinChunkInput, chunk_doc_with_builtin_splitter

        activity_input = BuiltinChunkInput(
            strategy_id=strategy_id,
            splitter="sentence",
            docs_root=str(tmp_path),
            doc_id=doc_id,
            content_hash=content_hash,
        )

        env = ActivityEnvironment()
        await env.run(chunk_doc_with_builtin_splitter, activity_input)

        loaded = load_chunk_cache(strategy_id)
        assert content_hash in loaded
        got = loaded[content_hash]
        assert len(got) >= 1
        # All chunks should be tagged with strategy and source_doc (LlamaIndex baseline path).
        for node in got:
            assert node.metadata.get("strategy") == strategy_id
            assert node.metadata.get("source_doc") == doc_id
        # Reassembled text should match original (minus possible whitespace normalization).
        reassembled = " ".join(n.text.strip() for n in got)
        assert "First sentence" in reassembled and "Third one" in reassembled
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)


@pytest.mark.asyncio
async def test_builtin_splitter_on_txt_and_docx_tree(tmp_path):
    """
    Built-in splitter activity on a small directory tree with .txt and .docx.
    Verifies the worker uses extract_file_content (not raw read_text) so PDF/DOCX work.
    """
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        # Simulate docs tree: subdir with one .txt and one .docx
        sub = tmp_path / "subdir"
        sub.mkdir()
        txt_path = sub / "doc.txt"
        txt_path.write_text("TXT content. Second sentence here.", encoding="utf-8")
        docx_path = sub / "doc.docx"
        try:
            import docx  # type: ignore[import]
            d = docx.Document()
            d.add_paragraph("DOCX content. Another sentence in Word.")
            d.save(str(docx_path))
        except ImportError:
            pytest.skip("python-docx required for DOCX test")

        doc_ids = ["subdir/doc.txt", "subdir/doc.docx"]
        from temporal_activities import (
            LoadDocManifestInput,
            BuiltinChunkInput,
            chunk_doc_with_builtin_splitter,
            load_doc_manifest,
        )

        env = ActivityEnvironment()
        manifest = await env.run(
            load_doc_manifest,
            LoadDocManifestInput(docs_root=str(tmp_path), doc_ids=doc_ids),
        )
        assert len(manifest) == 2, "manifest should include both .txt and .docx (extract_file_content used)"
        by_id = {m["doc_id"]: m["content_hash"] for m in manifest}
        assert "subdir/doc.txt" in by_id
        assert "subdir/doc.docx" in by_id

        strategy_id = "test_builtin_txt_docx"
        for m in manifest:
            await env.run(
                chunk_doc_with_builtin_splitter,
                BuiltinChunkInput(
                    strategy_id=strategy_id,
                    splitter="sentence",
                    docs_root=str(tmp_path),
                    doc_id=m["doc_id"],
                    content_hash=m["content_hash"],
                ),
            )

        loaded = load_chunk_cache(strategy_id)
        assert len(loaded) == 2
        texts = set()
        for nodes in loaded.values():
            for n in nodes:
                texts.add(n.text.strip())
        assert any("TXT content" in t for t in texts)
        assert any("DOCX content" in t or "Word" in t for t in texts)
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)


@pytest.mark.skipif(
    not os.environ.get("OPENAI_API_KEY"),
    reason="Set OPENAI_API_KEY to run real LLM chunking test (run manually)",
)
@pytest.mark.asyncio
async def test_real_llm_chunking_in_activity(tmp_path):
    """
    Real LLM chunking inside a Temporal activity (no mock). Requires OPENAI_API_KEY.
    Run manually: OPENAI_API_KEY=sk-... pytest tests/test_activity_logic.py::test_real_llm_chunking_in_activity -v
    """
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        from llama_index.core import Settings
        from llama_index.llms.openai import OpenAI

        Settings.llm = OpenAI(model="gpt-4o-mini", temperature=0.1)

        doc_content = (
            "The company was founded in 2020. It builds developer tools. "
            "The flagship product is a chunking library for RAG."
        )
        doc_file = tmp_path / "real_llm_doc.txt"
        doc_file.write_text(doc_content, encoding="utf-8")
        doc_id = "real_llm_doc.txt"
        strategy_id = "test_real_llm"
        content_hash = compute_content_hash(doc_content)

        from temporal_activities import ChunkDocInput, chunk_doc_with_strategy

        activity_input = ChunkDocInput(
            strategy_id=strategy_id,
            strategy_instruction="Split by sentence. One chunk per sentence.",
            model=None,
            docs_root=str(tmp_path),
            doc_id=doc_id,
            content_hash=content_hash,
        )

        env = ActivityEnvironment()
        await env.run(chunk_doc_with_strategy, activity_input)

        loaded = load_chunk_cache(strategy_id)
        assert content_hash in loaded
        got = loaded[content_hash]
        assert len(got) >= 1
        for node in got:
            assert node.metadata.get("strategy") == strategy_id
            assert node.metadata.get("source_doc") == doc_id
        all_text = " ".join(n.text for n in got)
        assert "founded in 2020" in all_text or "developer tools" in all_text or "chunking" in all_text
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)


# --- Phase 2: load_doc_manifest and get_cached_hashes activities ---


@pytest.mark.asyncio
async def test_load_doc_manifest_and_get_cached_hashes(tmp_path):
    """Phase 2: load_doc_manifest returns doc_id+content_hash; get_cached_hashes returns hashes in cache."""
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        (tmp_path / "a.txt").write_text("Alpha", encoding="utf-8")
        (tmp_path / "b.txt").write_text("Beta", encoding="utf-8")
        from temporal_activities import (
            GetCachedHashesInput,
            LoadDocManifestInput,
            get_cached_hashes,
            load_doc_manifest,
        )

        env = ActivityEnvironment()
        manifest = await env.run(
            load_doc_manifest,
            LoadDocManifestInput(docs_root=str(tmp_path), doc_ids=["a.txt", "b.txt"]),
        )
        assert len(manifest) == 2
        doc_ids = {m["doc_id"] for m in manifest}
        assert doc_ids == {"a.txt", "b.txt"}
        hash_alpha = compute_content_hash("Alpha")
        hash_beta = compute_content_hash("Beta")
        hashes = {m["content_hash"] for m in manifest}
        assert hashes == {hash_alpha, hash_beta}

        # Empty cache -> get_cached_hashes returns []
        cached = await env.run(get_cached_hashes, GetCachedHashesInput(strategy_id="s_phase2"))
        assert cached == []

        # After appending one doc's chunks, get_cached_hashes returns that hash
        append_chunk_cache("s_phase2", hash_alpha, [TextNode(text="x", metadata={})])
        cached = await env.run(get_cached_hashes, GetCachedHashesInput(strategy_id="s_phase2"))
        assert hash_alpha in cached
        assert hash_beta not in cached
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)


# --- Step 1.5: Verify retries (workflow + worker; one doc fails once, retry succeeds) ---

TASK_QUEUE_RETRY = "rechunk-strategy-chunking"


@pytest.mark.asyncio
async def test_activity_retry_one_doc_fails_then_succeeds(tmp_path):
    """
    Step 1.5: Run full workflow with 2 docs; second doc's activity fails on first attempt
    (simulated rate-limit). Temporal retries; on retry the activity succeeds. Assert both
    docs end up in cache. Uses WorkflowEnvironment (may download Temporal test server).
    """
    os.environ["RECHUNK_STRATEGY_CACHE_DIR"] = str(tmp_path)
    try:
        # Two docs so we can assert both complete (one after retry).
        (tmp_path / "ok.txt").write_text("Doc A content.", encoding="utf-8")
        (tmp_path / "fail_once.txt").write_text("Doc B content.", encoding="utf-8")
        hash_a = compute_content_hash("Doc A content.")
        hash_b = compute_content_hash("Doc B content.")

        from temporalio import activity
        from temporalio.exceptions import ApplicationError
        from temporalio.worker import Worker
        from temporalio.testing import WorkflowEnvironment

        from temporal_activities import (
            ChunkDocInput,
            chunk_doc_with_strategy as real_chunk_doc_with_strategy,
            get_cached_hashes,
            load_doc_manifest,
            log_workflow_summary,
        )
        from temporal_workflows import StrategyChunkingWorkflow, StrategyChunkingInput

        # Wrapper: fail on first attempt only for fail_once.txt (simulated rate-limit).
        @activity.defn(name="chunk_doc_with_strategy")
        async def chunk_doc_fail_once_then_ok(input: ChunkDocInput):
            info = activity.info()
            if input.doc_id == "fail_once.txt" and info.attempt == 1:
                raise ApplicationError(
                    "Simulated rate limit for retry test",
                    non_retryable=False,
                )
            return await real_chunk_doc_with_strategy(input)

        # Mock LLM so real activity does not call OpenAI.
        mock_nodes_a = [
            TextNode(text="Doc A chunk", metadata={"strategy": "test_retry", "source_doc": "ok.txt"}),
        ]
        mock_nodes_b = [
            TextNode(text="Doc B chunk", metadata={"strategy": "test_retry", "source_doc": "fail_once.txt"}),
        ]

        def mock_get_nodes(docs):
            if not docs:
                return []
            # Identify doc by content (activity runs one doc at a time).
            if docs[0].text.strip() == "Doc A content.":
                return mock_nodes_a
            return mock_nodes_b

        with patch(
            "rechunk.node_parser.LLMNodeParser.get_nodes_from_documents",
            side_effect=mock_get_nodes,
        ):
            try:
                env = await WorkflowEnvironment.start_time_skipping()
            except RuntimeError as e:
                if "test server" in str(e).lower() or "Failed starting" in str(e):
                    pytest.skip(
                        "Temporal test server unavailable (e.g. no network). "
                        "Run with network to verify retries."
                    )
                raise
            async with env:
                worker = Worker(
                    env.client,
                    task_queue=TASK_QUEUE_RETRY,
                    workflows=[StrategyChunkingWorkflow],
                    activities=[
                        chunk_doc_fail_once_then_ok,
                        load_doc_manifest,
                        get_cached_hashes,
                        log_workflow_summary,
                    ],
                )
                worker_task = asyncio.create_task(worker.run())
                try:
                    workflow_input = StrategyChunkingInput(
                        strategy_id="test_retry",
                        kind="llm",
                        docs_root=str(tmp_path),
                        doc_ids=["ok.txt", "fail_once.txt"],
                        strategy_instruction="Split.",
                        model=None,
                        splitter="sentence",
                    )
                    handle = await env.client.start_workflow(
                        StrategyChunkingWorkflow.run,
                        workflow_input,
                        id="test-retry-1",
                        task_queue=TASK_QUEUE_RETRY,
                    )
                    await handle.result()
                    # Both docs should be in cache (second after retry).
                    loaded = load_chunk_cache("test_retry")
                    assert hash_a in loaded, "First doc should be in cache"
                    assert hash_b in loaded, "Second doc should be in cache after retry"
                    assert len(loaded[hash_a]) >= 1
                    assert len(loaded[hash_b]) >= 1
                finally:
                    worker_task.cancel()
                    try:
                        await worker_task
                    except asyncio.CancelledError:
                        pass
    finally:
        os.environ.pop("RECHUNK_STRATEGY_CACHE_DIR", None)
